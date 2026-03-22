from docx import Document
import sys
import os

def read_doc_to_file(filepath, outpath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    try:
        doc = Document(filepath)
        with open(outpath, 'w', encoding='utf-8') as f:
            f.write(f"--- CONTENT OF {filepath} ---\n")
            for para in doc.paragraphs:
                if para.text.strip():
                    f.write(para.text + "\n")
            f.write("\n\n")
            
            f.write(f"--- TABLES IN {filepath} ---\n")
            for table in doc.tables:
                for row in table.rows:
                    f.write(" | ".join([cell.text.strip() for cell in row.cells]) + "\n")
            f.write("\n\n")
    except Exception as e:
        print(f"Error reading {filepath}: {e}")

read_doc_to_file("Ecolab - Bang de xuat hop tac.docx", "ecolab_vn.log")
read_doc_to_file("Ecolab proposal.docx", "ecolab_en.log")
print("Done writing to logs.")
